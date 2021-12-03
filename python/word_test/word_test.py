from docx import Document

doc = Document('test.docx')

for p in doc.paragraphs:
    print(f"{p.text} ({p.style.name})")

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                print(f"{p.text} ({p.style.name})")

# Iterate in order
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

for o in doc._element.body:
    if isinstance(o, CT_P):
        p = Paragraph(o, doc)
        print(f"{p.text} ({p.style.name})")
    elif isinstance(o, CT_Tbl):
        t = Table(o, doc)
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    print(f"{p.text} ({p.style.name})")

# 00h05 : yes!
